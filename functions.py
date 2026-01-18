import streamlit as st
import os
import time
from pathlib import Path
# import wave
# import pyaudio
from pydub import AudioSegment
from audiorecorder import audiorecorder
import re
import numpy as np
from scipy.io.wavfile import write
from langchain.prompts import (
    ChatPromptTemplate,
    HumanMessagePromptTemplate,
    MessagesPlaceholder,
)
from langchain.schema import SystemMessage
from langchain.memory import ConversationSummaryBufferMemory
from langchain_openai import ChatOpenAI
from langchain.chains import ConversationChain
import constants as ct
import uuid  # 一意のファイル名生成のためにuuidをインポート
import docx
import datetime

def record_audio(audio_input_file_path):
    """
    音声入力を受け取って音声ファイルを作成
    """

    audio = audiorecorder(
        start_prompt="発話開始",
        pause_prompt="やり直す",
        stop_prompt="発話終了",
        start_style={"color":"white", "background-color":"black"},
        pause_style={"color":"gray", "background-color":"white"},
        stop_style={"color":"white", "background-color":"black"}
    )

    if len(audio) > 0:
        audio.export(audio_input_file_path, format="wav")
    else:
        st.stop()

def transcribe_audio(audio_input_file_path):
    """
    音声入力ファイルから文字起こしテキストを取得
    Args:
        audio_input_file_path: 音声入力ファイルのパス
    Returns:
        transcript: Whisperの文字起こし結果
        warning_message: 警告メッセージ（なければNone）
    """

    # 音声ファイルの長さをチェック
    audio = AudioSegment.from_wav(audio_input_file_path)
    duration_seconds = len(audio) / 1000.0  # ミリ秒から秒に変換
    
    warning_message = None
    
    # 音声の長さをチェック
    if duration_seconds < 0.5:
        warning_message = "⚠️ 音声が非常に短いです。もう一度、発話してみてください。"

    with open(audio_input_file_path, 'rb') as audio_input_file:
        transcript = st.session_state.openai_obj.audio.transcriptions.create(
            model="whisper-1",
            file=audio_input_file,
            language="en"
        )
    
    # 文字起こし結果が空または非常に短い場合
    if not transcript.text or len(transcript.text.strip()) < 3:
        warning_message = "⚠️ 音声を認識できませんでした。もう一度、はっきりと発話してみてください。"
    
    # 音声入力ファイルを削除
    os.remove(audio_input_file_path)

    return transcript, warning_message


def transcribe_and_handle_language(audio_input_file_path):
    """
    音声を自動判定で文字起こしし、もし英語でない（日本語と判断された）場合は日本語で再文字起こしして
    日本語→英訳、さらに英訳文のカタカナ発音を生成するワークフローを提供する。

    Returns:
        result: dict with keys:
            detected_language: 'en' or 'ja' or 'unknown'
            original_text: transcription text (in detected language)
            translated_english: (if detected ja) English translation string else None
            katakana_pron: (if translated) Katakana pronunciation string prefixed with '発音：' else None
            warning_message: warning or None
    """

    warning_message = None

    def _extract_text(obj):
        if obj is None:
            return ''
        if isinstance(obj, dict):
            return obj.get('text', '') or ''
        if hasattr(obj, 'text'):
            return getattr(obj, 'text') or ''
        try:
            # fallback to string conversion
            return str(obj)
        except Exception:
            return ''

    def _extract_language(obj):
        if obj is None:
            return None
        if isinstance(obj, dict) and 'language' in obj:
            return obj.get('language')
        if hasattr(obj, 'language'):
            return getattr(obj, 'language')
        return None

    # まず自動判定で文字起こし（言語パラメータを送らない）
    try:
        with open(audio_input_file_path, 'rb') as audio_input_file:
            transcript_auto = st.session_state.openai_obj.audio.transcriptions.create(
                model="whisper-1",
                file=audio_input_file
            )
    except Exception as e:
        return {
            'detected_language': 'unknown',
            'original_text': '',
            'translated_english': None,
            'katakana_pron': None,
            'warning_message': f"音声の文字起こしに失敗しました: {e}"
        }

    text_auto = _extract_text(transcript_auto)
    detected = _extract_language(transcript_auto)

    # 判定ルール拡張:
    # - もし既にAPIが言語を返していればそれを使う
    # - 返ってこなければテキストを解析して判定する
    #   * 英単語が3つ以上含まれていれば英語と扱う（和製英語対応）
    #   * それ以外で日本語文字が含まれていれば日本語
    #   * それ以外は英語とする
    # Count English-like tokens regardless of whisper's detected field.
    eng_words = re.findall(r"[A-Za-z]+(?:'[A-Za-z]+)?", text_auto)
    # Rule: If 3 or more English tokens exist, treat as English (override Japanese detection).
    if len(eng_words) >= 3:
        detected = 'en'
    else:
        # If whisper did not supply language, fall back to Japanese character detection
        if not detected:
            if re.search('[\u3040-\u30ff\u4e00-\u9fff]', text_auto):
                detected = 'ja'
            else:
                detected = 'en'

    result = {
        'detected_language': detected,
        'original_text': text_auto,
        'translated_english': None,
        'katakana_pron': None,
        'warning_message': warning_message
    }

    # 日本語と判定された場合は日本語で再度文字起こしして正確な日本語テキストを得る
    if str(detected).lower().startswith('ja'):
        try:
            with open(audio_input_file_path, 'rb') as audio_input_file:
                transcript_ja = st.session_state.openai_obj.audio.transcriptions.create(
                    model="whisper-1",
                    file=audio_input_file,
                    language="ja"
                )
            text_ja = _extract_text(transcript_ja)
            result['original_text'] = text_ja

            # 日本語を英訳（LLMを利用）
            try:
                translate_prompt = f"""
Translate the following Japanese sentence to natural, concise English. Provide only the English translation.

Japanese: "{text_ja}"

English:
"""
                translated_en = st.session_state.llm.predict(translate_prompt).strip()
                result['translated_english'] = translated_en
            except Exception as e:
                result['warning_message'] = f"日本語の英訳に失敗しました: {e}"

            # カタカナ発音を生成（LLMに依頼）
            if result.get('translated_english'):
                try:
                    kana_prompt = f"""
Katakana phonetic transcription of the following English sentence, optimized for being read aloud by a Japanese speaker.

Apply the following rules strictly:

1. Reduction:
   - Identify the syllables that are barely pronounced or reduced in natural spoken English.
   - Enclose ONLY the reduced/weak syllables in parentheses (), not the main stressed part.
   - Parentheses indicate a weak, reduced sound that is still pronounced, not a completely omitted sound.
   - Representative examples:
     - /t/ final: “efficient” → エフィシェン(ト)
     - /d/ final: “good” → グ(ドゥ)
     - /g/ final (nasalized): “doing” → ドゥイ(ン)
     - /k/ final (unreleased): “back” → バ(ク)

2. Liaison / Linking:
   - Apply natural connected-speech transformations aggressively.
   - Merge consonant → vowel connections.
   - Examples:
     - “want to” → ワナ
     - “going to” → ゴナ
     - “look at it” → ルカリッ

3. Flapping:
   - Apply American flapping for /t/ and /d/ between vowels.
   - Examples:
     - “water” → ワラー
     - “get it” → ゲリッ
     - “better” → ベラー

Output:
- Katakana only.
- Single line.
- No English, no explanations.

English: "{result['translated_english']}"
"""
                    katakana = st.session_state.llm.predict(kana_prompt).strip()
                    if katakana:
                        if not katakana.startswith('発音：'):
                            katakana = f"発音：{katakana}"
                        result['katakana_pron'] = katakana
                except Exception as e:
                    # 非致命: katakana生成失敗
                    result['warning_message'] = (result.get('warning_message') or '') + f" カタカナ生成に失敗しました: {e}"

        except Exception as e:
            result['warning_message'] = f"日本語での再文字起こしに失敗しました: {e}"

    # 音声入力ファイルを削除
    try:
        os.remove(audio_input_file_path)
    except Exception:
        pass

    return result

def save_to_wav(llm_response_audio, audio_output_file_path):
    """
    一旦mp3形式で音声ファイル作成後、wav形式に変換
    Args:
        llm_response_audio: LLMからの回答の音声データ
        audio_output_file_path: 出力先のファイルパス
    """

    temp_audio_output_filename = f"{ct.AUDIO_OUTPUT_DIR}/temp_audio_output_{int(time.time())}.mp3"
    with open(temp_audio_output_filename, "wb") as temp_audio_output_file:
        temp_audio_output_file.write(llm_response_audio)
    
    audio_mp3 = AudioSegment.from_file(temp_audio_output_filename, format="mp3")
    audio_mp3.export(audio_output_file_path, format="wav")

    # 音声出力用に一時的に作ったmp3ファイルを削除
    os.remove(temp_audio_output_filename)

# def play_wav(audio_output_file_path, speed=1.0):
#     """
#     音声ファイルの読み上げ
#     Args:
#         audio_output_file_path: 音声ファイルのパス
#         speed: 再生速度（1.0が通常速度、0.5で半分の速さ、2.0で倍速など）
#     """

#     # 音声ファイルの読み込み
#     audio = AudioSegment.from_wav(audio_output_file_path)
    
#     # 速度を変更
#     if speed != 1.0:
#         # frame_rateを変更することで速度を調整
#         modified_audio = audio._spawn(
#             audio.raw_data, 
#             overrides={"frame_rate": int(audio.frame_rate * speed)}
#         )
#         # 元のframe_rateに戻すことで正常再生させる（ピッチを保持したまま速度だけ変更）
#         modified_audio = modified_audio.set_frame_rate(audio.frame_rate)

#         modified_audio.export(audio_output_file_path, format="wav")

#     # PyAudioで再生
#     with wave.open(audio_output_file_path, 'rb') as play_target_file:
#         p = pyaudio.PyAudio()
#         stream = p.open(
#             format=p.get_format_from_width(play_target_file.getsampwidth()),
#             channels=play_target_file.getnchannels(),
#             rate=play_target_file.getframerate(),
#             output=True
#         )

#         data = play_target_file.readframes(1024)
#         while data:
#             stream.write(data)
#             data = play_target_file.readframes(1024)

#         stream.stop_stream()
#         stream.close()
#         p.terminate()
    
#     # LLMからの回答の音声ファイルを削除
#     os.remove(audio_output_file_path)

# ！上記関数の代替として、再生速度変更用関数を追加
def change_speed(input_wav, output_wav, speed):
    """
    音声ファイルの再生速度を変更して保存
    Args:
        input_wav: 入力WAVファイルのパス
        output_wav: 出力WAVファイルのパス
        speed: 再生速度（1.0が通常速度、0.5で半分の速さ、2.0で倍速など）
    """
    audio = AudioSegment.from_wav(input_wav)
    audio = audio.speedup(playback_speed=speed)
    audio.export(output_wav, format="wav")

def create_chain(system_template):
    """
    LLMによる回答生成用のChain作成
    """

    prompt = ChatPromptTemplate.from_messages([
        SystemMessage(content=system_template),
        MessagesPlaceholder(variable_name="history"),
        HumanMessagePromptTemplate.from_template("{input}")
    ])
    chain = ConversationChain(
        llm=st.session_state.llm,
        memory=st.session_state.memory,
        prompt=prompt
    )

    return chain

def create_problem_and_play_audio():
    """
    問題生成と音声ファイルの再生
    Args:
        chain: 問題文生成用のChain
        speed: 再生速度（1.0が通常速度、0.5で半分の速さ、2.0で倍速など）
        openai_obj: OpenAIのオブジェクト
    """

    # 問題文を生成するChainを実行し、問題文を取得
    problem = st.session_state.chain_create_problem.predict(input="")

    # LLMからの回答を音声データに変換
    llm_response_audio = st.session_state.openai_obj.audio.speech.create(
        model="tts-1",
        voice="alloy",
        input=problem
    )

    # 音声ファイルの作成
    audio_output_file_path = f"{ct.AUDIO_OUTPUT_DIR}/audio_output_{int(time.time())}.wav"
    save_to_wav(llm_response_audio.content, audio_output_file_path)

    # 音声ファイルの再生処理を変更
    # play_wav(audio_output_file_path, st.session_state.speed)  # 元のコードをコメントアウト
    if st.session_state.speed != 1.0:
        temp_audio_path = f"temp_{uuid.uuid4().hex}.wav"
        change_speed(audio_output_file_path, temp_audio_path, st.session_state.speed)
        audio_output_file_path = temp_audio_path

    return problem, llm_response_audio, audio_output_file_path

def create_evaluation():
    """
    ユーザー入力値の評価生成
    """

    llm_response_evaluation = st.session_state.chain_evaluation.predict(input="")

    return llm_response_evaluation

def correct_user_input(user_text, level):
    """
    ユーザーの英語発話を添削し、より良い表現を提示
    Args:
        user_text: ユーザーの英語発話
        level: ユーザーの英語レベル
    Returns:
        correction: 添削結果（改善が必要ない場合はNone）
    """
    correction_prompt = f"""
You are an English grammar expert. Analyze the following English sentence and provide corrections if needed.

User's English Level: {level}
User's sentence: "{user_text}"

If the sentence has grammatical errors or could be improved:
1. Provide a corrected/improved version
2. Briefly explain the improvements in Japanese
3. Keep the original meaning intact

If the sentence is already correct and natural, simply respond with: "Perfect! No corrections needed."

Format your response as:
【改善案】
[Improved English sentence]

【解説】
[Brief explanation in Japanese]
"""
    
    correction = st.session_state.llm.predict(correction_prompt)
    
    # 添削が不要な場合はNoneを返す
    if "Perfect" in correction or "No corrections needed" in correction:
        return None
    
    return correction

def translate_to_japanese(english_text):
    """
    英語テキストを日本語に翻訳
    Args:
        english_text: 英語テキスト
    Returns:
        japanese_text: 日本語訳
    """
    translation_prompt = f"""
Translate the following English text to natural Japanese.
Provide only the Japanese translation without any additional explanation.

English: "{english_text}"

Japanese:
"""
    
    japanese_text = st.session_state.llm.predict(translation_prompt)
    
    return japanese_text.strip()

def correct_and_translate_batch(user_text, ai_response, level):
    """
    ユーザー発話の添削とAI返事の翻訳を1回のLLM呼び出しで取得（トークン節約）
    Args:
        user_text: ユーザーの英語発話
        ai_response: AIの英語返事
        level: ユーザーの英語レベル
    Returns:
        correction: 添削結果（改善不要の場合None）
        translation: 日本語訳
    """
    batch_prompt = f"""You must perform TWO tasks and return results in EXACT format below.

TASK 1 - Grammar Check
User's Level: {level}
User said: "{user_text}"

If there are errors or improvements needed:
- Provide corrected English sentence
- Explain improvements in Japanese
If already perfect, output only: PERFECT

TASK 2 - Translation
AI said: "{ai_response}"
Translate to natural Japanese.

CRITICAL: Use EXACTLY this format with markers:
<<<CORRECTION_START>>>
[Your correction result here - either "PERFECT" or corrected sentence with Japanese explanation]
<<<CORRECTION_END>>>

<<<TRANSLATION_START>>>
[Japanese translation here]
<<<TRANSLATION_END>>>
"""
    
    result = st.session_state.llm.predict(batch_prompt)
    
    # 結果を分割（改善されたパース処理）
    correction = None
    translation = ""
    
    try:
        # 添削部分を抽出
        if "<<<CORRECTION_START>>>" in result and "<<<CORRECTION_END>>>" in result:
            correction_start = result.find("<<<CORRECTION_START>>>") + len("<<<CORRECTION_START>>>")
            correction_end = result.find("<<<CORRECTION_END>>>")
            correction_part = result[correction_start:correction_end].strip()
            
            # "PERFECT"でなければ添削結果として保持
            if "PERFECT" not in correction_part.upper():
                correction = correction_part
        
        # 翻訳部分を抽出
        if "<<<TRANSLATION_START>>>" in result and "<<<TRANSLATION_END>>>" in result:
            translation_start = result.find("<<<TRANSLATION_START>>>") + len("<<<TRANSLATION_START>>>")
            translation_end = result.find("<<<TRANSLATION_END>>>")
            translation = result[translation_start:translation_end].strip()
    
    except Exception as e:
        # パース失敗時のフォールバック
        st.warning(f"⚠️ 添削・翻訳の解析に失敗しました: {e}")
        # 少なくとも何か返す
        if "PERFECT" not in result.upper():
            correction = result[:500]  # 最初の部分を返す
    
    return correction, translation


def _split_blocks_from_docx(filepath):
    """
    Read a .docx file and split into blocks separated by empty paragraphs.
    Each block is returned as a single string (preserving newlines).
    """
    try:
        doc = docx.Document(filepath)
    except Exception:
        return []
    blocks = []
    cur_lines = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text == "":
            if cur_lines:
                blocks.append("\n".join(cur_lines).strip())
                cur_lines = []
        else:
            cur_lines.append(text)
    if cur_lines:
        blocks.append("\n".join(cur_lines).strip())
    return blocks


def parse_block_to_parts(block_text):
    """
    Expect block to contain lines in order: English, Japanese, Katakana (発音), Grammar unit.
    Be permissive: return dict with keys 'english','japanese','katakana','grammar_unit'.
    """
    lines = [ln.strip() for ln in block_text.splitlines() if ln.strip()]
    parts = {'english': '', 'japanese': '', 'katakana': '', 'grammar_unit': ''}
    if not lines:
        return parts
    parts['english'] = lines[0]
    if len(lines) > 1:
        parts['japanese'] = lines[1]
    if len(lines) > 2:
        parts['katakana'] = lines[2]
    if len(lines) > 3:
        parts['grammar_unit'] = lines[3]
    # If more lines exist, append them to grammar_unit
    if len(lines) > 4:
        parts['grammar_unit'] += '\n' + '\n'.join(lines[4:])
    return parts


def pick_shadowing_block_from_documents(documents_dir):
    """
    Choose one of the document files based on day parity and select a block
    using the sum of the digits of today's date as index.

    Returns: dict from parse_block_to_parts or None if failure.
    """
    try:
        p = Path(documents_dir)
        files = sorted([x for x in p.glob('*.docx')])
        if not files:
            return None

        # If cached file list differs from current, reload cache but try to preserve selection
        current_file_names = [str(x) for x in files]
        cached_files = st.session_state.get('shadowing_files')
        prev_selected = st.session_state.get('shadowing_selected_file')
        if cached_files != current_file_names:
            # reload all files into cache
            st.session_state['shadowing_files'] = current_file_names
            st.session_state['shadowing_rows'] = {}
            for f in files:
                rows = []
                try:
                    doc = docx.Document(f)
                    for table in doc.tables:
                        for r in table.rows:
                            cells = [c.text.strip() for c in r.cells]
                            if not any(cells):
                                continue
                            while len(cells) < 4:
                                cells.append('')
                            rows.append({'english': cells[0], 'japanese': cells[1], 'katakana': cells[2], 'grammar_unit': cells[3]})
                except Exception:
                    # fallback to paragraph blocks
                    blocks = _split_blocks_from_docx(str(f))
                    for b in blocks:
                        parts = parse_block_to_parts(b)
                        rows.append(parts)
                st.session_state['shadowing_rows'][str(f)] = rows

            # determine chosen file: prefer previous selection if still available
            today = datetime.date.today()
            day = today.day
            if prev_selected and prev_selected in current_file_names:
                chosen_file_str = prev_selected
            else:
                file_index = day % len(files)
                chosen_file_str = str(files[file_index])

            # initial row index based on date digit sum (if not already set)
            ymd = today.strftime('%Y%m%d')
            s = sum(int(ch) for ch in ymd if ch.isdigit())
            rows_for_file = st.session_state['shadowing_rows'].get(chosen_file_str, [])
            if not rows_for_file:
                st.session_state['shadowing_selected_file'] = chosen_file_str
                st.session_state['shadowing_row_index'] = 0
            else:
                # if previously had an index and still valid, keep it; otherwise initialize
                existing_idx = st.session_state.get('shadowing_row_index')
                if existing_idx is None or existing_idx >= len(rows_for_file):
                    st.session_state['shadowing_selected_file'] = chosen_file_str
                    st.session_state['shadowing_row_index'] = s % len(rows_for_file)
                else:
                    st.session_state['shadowing_selected_file'] = chosen_file_str
                    # keep existing_idx as-is

        # Use cached selection
        chosen_file_str = st.session_state.get('shadowing_selected_file')
        rows = st.session_state.get('shadowing_rows', {}).get(chosen_file_str, [])
        if not rows:
            return None

        # get current index and then increment for next time
        idx = st.session_state.get('shadowing_row_index', 0)
        selected = rows[idx]
        selected['_source_file'] = Path(chosen_file_str).name
        selected['_block_index'] = idx

        # increment index for next invocation
        st.session_state['shadowing_row_index'] = (idx + 1) % len(rows)

        return selected
    except Exception:
        return None


def explain_grammar_for_shadowing(english_sentence, grammar_unit=''):
    """
    Use LLM to produce a concise grammar explanation in Japanese for the given sentence.
    Keep it short (2-4 sentences), focused on the grammar point and one short example.
    """
    if not hasattr(st.session_state, 'llm'):
        return '（文法解説は利用できません：LLMが初期化されていません）'

    prompt = f"""
あなたは日本語で教える英語講師です。以下の英文と対応する文法項目について、学習者向けに簡潔に解説してください。

英文: {english_sentence}
文法項目（参考）: {grammar_unit}

要求:
- 日本語で2〜4文の簡潔な解説
- 使い方のポイントと注意点を一つ示す
- 例文を1つ短く示す（英語と日本語訳）
"""
    try:
        res = st.session_state.llm.predict(prompt)
        return res
    except Exception as e:
        return f'文法解説の生成に失敗しました: {e}'